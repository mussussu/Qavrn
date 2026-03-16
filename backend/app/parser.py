from __future__ import annotations

import hashlib
import json as _json
import logging
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any, Dict

logger = logging.getLogger(__name__)

# Extensions that are read verbatim as plain UTF-8 text
_PLAIN_EXTS = {".txt", ".csv", ".rst", ".log", ".yaml", ".yml", ".toml", ".env"}

# Source-code extensions — read as plain text with a header line
_CODE_EXTS = {
    ".py", ".js", ".ts", ".java", ".cpp", ".c",
    ".rs", ".go", ".rb", ".php", ".swift", ".kt",
}


@dataclass
class Document:
    file_path: str
    filename: str
    content: str
    file_type: str
    last_modified: float
    metadata: Dict[str, Any] = field(default_factory=dict)

    @property
    def document_id(self) -> str:
        """Stable hash of the absolute file path."""
        return hashlib.sha256(self.file_path.encode()).hexdigest()


class DocumentParser:
    """Detects file type by extension and extracts plain text."""

    def parse(self, file_path: str | Path) -> Document:
        path = Path(file_path).resolve()
        ext = path.suffix.lower()
        last_modified = path.stat().st_mtime

        content = self._extract(path, ext)

        return Document(
            file_path=str(path),
            filename=path.name,
            content=content,
            file_type=ext.lstrip("."),
            last_modified=last_modified,
            metadata={
                "file_size": path.stat().st_size,
                "extension": ext,
            },
        )

    # ------------------------------------------------------------------
    # Private dispatch
    # ------------------------------------------------------------------

    def _extract(self, path: Path, ext: str) -> str:
        # Named handlers take priority
        named: Dict[str, Any] = {
            ".pdf":  self._pdf,
            ".docx": self._docx,
            ".md":   self._markdown,
            ".html": self._html,
            ".json": self._json,
            ".xml":  self._xml,
            ".eml":  self._eml,
            ".epub": self._epub,
        }
        if ext in named:
            return named[ext](path)
        if ext in _PLAIN_EXTS:
            return self._plain(path)
        if ext in _CODE_EXTS:
            return self._code(path)
        raise ValueError(f"Unsupported extension: {ext}")

    # ------------------------------------------------------------------
    # Original handlers
    # ------------------------------------------------------------------

    def _pdf(self, path: Path) -> str:
        try:
            import PyPDF2  # type: ignore
        except ImportError:
            raise ImportError("PyPDF2 is required for PDF parsing: pip install PyPDF2")

        text_parts: list[str] = []
        with open(path, "rb") as fh:
            reader = PyPDF2.PdfReader(fh)
            for page in reader.pages:
                try:
                    text_parts.append(page.extract_text() or "")
                except Exception as exc:
                    logger.warning("Could not extract text from PDF page in %s: %s", path, exc)
        return "\n".join(text_parts)

    def _docx(self, path: Path) -> str:
        try:
            import docx  # type: ignore
        except ImportError:
            raise ImportError("python-docx is required for DOCX parsing: pip install python-docx")

        doc = docx.Document(str(path))
        return "\n".join(para.text for para in doc.paragraphs)

    def _markdown(self, path: Path) -> str:
        try:
            import markdown  # type: ignore
            from html.parser import HTMLParser

            class _Stripper(HTMLParser):
                def __init__(self):
                    super().__init__()
                    self._parts: list[str] = []

                def handle_data(self, data: str) -> None:
                    self._parts.append(data)

                def get_text(self) -> str:
                    return "".join(self._parts)

            raw = path.read_text(encoding="utf-8", errors="replace")
            html = markdown.markdown(raw)
            stripper = _Stripper()
            stripper.feed(html)
            return stripper.get_text()
        except ImportError:
            logger.warning("markdown package not installed; returning raw MD text for %s", path)
            return path.read_text(encoding="utf-8", errors="replace")

    def _plain(self, path: Path) -> str:
        return path.read_text(encoding="utf-8", errors="replace")

    def _html(self, path: Path) -> str:
        from html.parser import HTMLParser

        class _Stripper(HTMLParser):
            def __init__(self):
                super().__init__()
                self._parts: list[str] = []
                self._skip_tags = {"script", "style"}
                self._current_skip: str | None = None

            def handle_starttag(self, tag: str, attrs) -> None:
                if tag in self._skip_tags:
                    self._current_skip = tag

            def handle_endtag(self, tag: str) -> None:
                if tag == self._current_skip:
                    self._current_skip = None

            def handle_data(self, data: str) -> None:
                if self._current_skip is None:
                    self._parts.append(data)

            def get_text(self) -> str:
                return " ".join(self._parts)

        raw = path.read_text(encoding="utf-8", errors="replace")
        stripper = _Stripper()
        stripper.feed(raw)
        return stripper.get_text()

    # ------------------------------------------------------------------
    # New handlers
    # ------------------------------------------------------------------

    def _code(self, path: Path) -> str:
        """Source-code files: prepend a header so the LLM has context."""
        content = path.read_text(encoding="utf-8", errors="replace")
        return f"Source code: {path.name}\n\n{content}"

    def _json(self, path: Path) -> str:
        """Pretty-print JSON so keys/values are readable as prose."""
        raw = path.read_text(encoding="utf-8", errors="replace")
        try:
            data = _json.loads(raw)
            return _json.dumps(data, indent=2, ensure_ascii=False)
        except _json.JSONDecodeError:
            return raw

    def _xml(self, path: Path) -> str:
        """Extract all text nodes from XML, stripping tags."""
        from xml.etree import ElementTree

        raw = path.read_text(encoding="utf-8", errors="replace")
        try:
            root = ElementTree.fromstring(raw)
            parts: list[str] = []
            for elem in root.iter():
                if elem.text and elem.text.strip():
                    parts.append(elem.text.strip())
                if elem.tail and elem.tail.strip():
                    parts.append(elem.tail.strip())
            return "\n".join(parts)
        except ElementTree.ParseError:
            return raw

    def _eml(self, path: Path) -> str:
        """Parse an RFC-822 email file and return headers + body text."""
        import email
        from email import policy as _policy

        raw = path.read_bytes()
        msg = email.message_from_bytes(raw, policy=_policy.default)

        parts: list[str] = [
            f"From: {msg.get('From', '')}",
            f"To: {msg.get('To', '')}",
            f"Subject: {msg.get('Subject', '')}",
            f"Date: {msg.get('Date', '')}",
            "",
        ]

        if msg.is_multipart():
            for part in msg.walk():
                if part.get_content_type() == "text/plain":
                    try:
                        parts.append(part.get_payload(decode=True).decode("utf-8", errors="replace"))
                    except Exception:
                        pass
        else:
            try:
                parts.append(msg.get_payload(decode=True).decode("utf-8", errors="replace"))
            except Exception:
                parts.append(str(msg.get_payload()))

        return "\n".join(parts)

    def _epub(self, path: Path) -> str:
        """Extract readable text from an EPUB ebook."""
        try:
            import ebooklib  # type: ignore
            from ebooklib import epub  # type: ignore
            from bs4 import BeautifulSoup  # type: ignore
        except ImportError:
            raise ImportError(
                "ebooklib and beautifulsoup4 are required for EPUB parsing: "
                "pip install ebooklib beautifulsoup4"
            )

        book = epub.read_epub(str(path), options={"ignore_ncx": True})
        text_parts: list[str] = []

        for item in book.get_items_of_type(ebooklib.ITEM_DOCUMENT):
            try:
                soup = BeautifulSoup(item.get_content(), "html.parser")
                text = soup.get_text(separator="\n", strip=True)
                if text:
                    text_parts.append(text)
            except Exception as exc:
                logger.warning("Could not extract EPUB chapter from %s: %s", path, exc)

        return "\n\n".join(text_parts)
